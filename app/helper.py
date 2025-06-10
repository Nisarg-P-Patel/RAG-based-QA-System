# app/helper.py

import os
import shutil
import csv
import datetime
from docx import Document
from models.rag import RAGPipeline
import gradio as gr

# === Directory Setup ===
FEEDBACK_DIR = "feedback"
EXPORTS_DIR = "exports"
FEEDBACK_FILE = os.path.join(FEEDBACK_DIR, "rag_feedback.csv")
VOTE_FILE = os.path.join(FEEDBACK_DIR, "rag_votes.csv")
EXPORT_PATH = os.path.join(EXPORTS_DIR, "rag_summary.docx")

os.makedirs(FEEDBACK_DIR, exist_ok=True)
os.makedirs(EXPORTS_DIR, exist_ok=True)

# === Session state for export ===
session_answers = []


def handle_query(query, state, summarize_docs, model_selection, pipeline_cache):
    rag = pipeline_cache[model_selection]["rag"]
    qa_chain = pipeline_cache[model_selection]["qa_chain"]

    print(f"[DEBUG] Query received: {query}")
    result = rag.run(query, summarize_docs=summarize_docs)

    answer = result['answer']
    context = result['context']
    category = result['category']
    sources = result['sources']
    confidence = result['confidence_score']
    highlighted_chunks = result['highlights']
    unique_docs = result['retrieved_chunks']

    state = {
        'query': query,
        'context': context,
        'category': category,
        'answer': answer,
        'docs': sources,
        'suggestion': "",
        'confidence': confidence,
        'model_used': model_selection
    }

    session_answers.append({
        'query': query,
        'category': category,
        'context': context,
        'answer': answer,
        'docs': sources,
        'confidence': confidence
    })

    display = f"""### ✅ Answer Generated

**Category:** {category}  
**Confidence Score:** {confidence}%

**Answer:**  
{answer}

**Top Sources:**  
""" + "\n".join([f"- {meta['source']}" for _, _, meta in unique_docs[:5]])

    chunk_display = "\n".join(highlighted_chunks)

    # return (
    #     display,
    #     state,
    #     chunk_display
    # )

    return (
        display,
        state,
        gr.update(visible=True),              # upvote_btn
        gr.update(visible=True),              # downvote_btn
        gr.update(visible=True),              # vote_ack
        gr.update(value = "", visible=True),  # suggestion textbox
        gr.update(visible=True),              # apply_btn
        gr.update(visible=True),              # final_display
        gr.update(value=chunk_display, visible=True)
    )


def apply_suggestion(suggestion, state, pipeline_cache):
    model = state.get("model_used", "tinyllama")
    qa_chain = pipeline_cache[model]["qa_chain"]

    if not suggestion.strip():
        final_answer = state.get('answer', 'No answer found.')
    else:
        improved_prompt = {
            "query": state['query'],
            "context": f"{state['context']}\n\nPrevious Answer: {state['answer']}\n\nSuggestion: {suggestion}",
            "category": state['category']
        }
        new_answer = qa_chain.invoke(improved_prompt)
        final_answer = new_answer['text'] if isinstance(new_answer, dict) else new_answer
        state['answer'] = final_answer

    session_answers.append({
        'query': state.get('query'),
        'category': state.get('category'),
        'context': state.get('context'),
        'answer': final_answer,
        'docs': state.get('docs'),
        'suggestion': suggestion,
        'confidence': state.get('confidence')
    })

    display = f"### ✅ Final Answer Stored\n\n**{final_answer}**\n\nYou may now ask another question or export all results as a Word document."

    return display, state


def generate_document():
    doc = Document()
    doc.add_heading("RAG QA Summary", level=0)

    for i, item in enumerate(session_answers, 1):
        doc.add_heading(f"Q{i}", level=1)
        table = doc.add_table(rows=5, cols=1)
        table.style = 'Table Grid'
        table.cell(0, 0).text = f"Query:\n{item['query']}"
        table.cell(1, 0).text = f"Category:\n{item['category']}"
        table.cell(2, 0).text = f"Final Answer:\n{item['answer']}"
        table.cell(3, 0).text = f"User Suggestion:\n{item.get('suggestion', '')}"
        table.cell(4, 0).text = f"Confidence Score:\n{item.get('confidence', 0)}%"

        doc.add_paragraph("")
        doc.add_heading("References", level=2)
        for j, doc_item in enumerate(item.get('docs', []), 1):
            source = doc_item.get('source', 'Unknown Source')
            doc.add_paragraph(f"{j}. {source}")
        doc.add_paragraph("=" * 50)

    doc.save(EXPORT_PATH)

    try:
        shutil.copy(EXPORT_PATH, "/content/drive/MyDrive/rag_summary.docx")
        print(f"[DEBUG] Document copied to Google Drive.")
    except Exception as e:
        print(f"[ERROR] Could not copy document to Drive: {e}")

    return EXPORT_PATH


def handle_feedback(feedback_text, state):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    feedback_data = {
        "timestamp": timestamp,
        "query": state.get("query", ""),
        "category": state.get("category", ""),
        "answer": state.get("answer", ""),
        "context": state.get("context", ""),
        "feedback": feedback_text,
        "sources": "; ".join([doc.get("source", "") for doc in state.get("docs", [])])
    }

    file_exists = os.path.exists(FEEDBACK_FILE)

    with open(FEEDBACK_FILE, "a", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=feedback_data.keys())
        if not file_exists:
            writer.writeheader()
        writer.writerow(feedback_data)

    try:
        shutil.copy(FEEDBACK_FILE, "/content/drive/MyDrive/rag_feedback.csv")
        print(f"[DEBUG] Feedback CSV backed up to Google Drive.")
    except Exception as e:
        print(f"[ERROR] Could not copy feedback to Drive: {e}")

    return "✅ Feedback submitted. Thank you!"


def handle_vote(vote_type, state):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    vote_data = {
        "timestamp": timestamp,
        "vote": vote_type,
        "query": state.get("query", ""),
        "category": state.get("category", ""),
        "answer": state.get("answer", ""),
        "context": state.get("context", ""),
        "sources": "; ".join([doc.get("source", "") for doc in state.get("docs", [])])
    }

    file_exists = os.path.exists(VOTE_FILE)

    with open(VOTE_FILE, "a", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=vote_data.keys())
        if not file_exists:
            writer.writeheader()
        writer.writerow(vote_data)

    try:
        shutil.copy(VOTE_FILE, "/content/drive/MyDrive/rag_votes.csv")
        print(f"[DEBUG] Vote CSV backed up to Google Drive.")
    except Exception as e:
        print(f"[ERROR] Could not copy vote to Drive: {e}")

    return f"✅ Your vote ({vote_type}) has been recorded. Thank you!"
